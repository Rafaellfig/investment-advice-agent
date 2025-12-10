"""Utilities for reading and writing files."""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Mm
from docx.oxml.ns import qn
from datetime import datetime
import locale
import io
import os
import re
locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")


def read_file(file_path: Path) -> str:
    """
    Reads the content of a text file.
    
    Args:
        file_path: Path of the file to be read
        
    Returns:
        File content as string
        
    Raises:
        FileNotFoundError: If the file is not found
        Exception: If another error occurs while reading the file
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {file_path}")
    except Exception as e:
        raise Exception(f"Error reading file {file_path}: {str(e)}")


def write_file(file_path: Path, content: str) -> None:
    """
    Writes content to a text file.
    
    Args:
        file_path: Path of the file to be written
        content: Content to be written to the file
        
    Raises:
        Exception: If an error occurs while writing the file
    """
    try:
        # Ensure the directory exists
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        raise Exception(f"Error writing file {file_path}: {str(e)}")


def set_document_margins(section, top_mm=25, bottom_mm=25, left_mm=25, right_mm=25):
    """Helper to configure document margins in millimeters."""
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)


def add_paragraph_with_bold(paragraph, text):
    """
    Adds text to a paragraph, processing **text** as bold formatting.
    
    Args:
        paragraph: The docx paragraph object to add text to.
        text: The text string that may contain **bold** markers.
    """
    # Pattern to match **text**
    pattern = r'\*\*(.*?)\*\*'
    
    # Split text by bold markers
    parts = re.split(pattern, text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        # Odd indices are the bold text (captured groups)
        if i % 2 == 1:
            run = paragraph.add_run(part)
            run.bold = True
        else:
            # Even indices are regular text
            paragraph.add_run(part)

def create_letter(
    output_file,
    image_path,
    sender_name,
    sender_title,
    city,
    recipient_name,
    subject_text,
    body_paragraphs,
    closing_line,
    signature_name,
    signature_title,
    chart_figure=None,
    top_mm=25,
    bottom_mm=25,
    left_mm=22,
    right_mm=22
):
    """
    Creates a formatted letter in a Word document using python-docx.

    Parameters:
        output_file (str): Path where the .docx will be saved.
        image_path (str): Path to the header image (logo).
        sender_name (str): Name of the sender.
        sender_title (str): Title/Role of the sender.
        city (str): City for the date line.
        recipient_name (str): Name of the receiver.
        subject_text (str): Subject title for the letter.
        body_paragraphs (list[str]): List of paragraphs for the body text.
        closing_line (str): Closing sentence before the signature.
        signature_name (str): Name to appear in the signature.
        signature_title (str): Title below signature.
        chart_figure (matplotlib.figure.Figure, optional): Matplotlib figure to insert in the document.
        top_mm, bottom_mm, left_mm, right_mm (int): Margin sizes.
    """
    doc = Document()

    # Default document font (ensures accent compatibility)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    # to ensure the font name is correctly applied in Word with languages:
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    font.size = Pt(11)

    # --- Margins ---
    section = doc.sections[0]
    set_document_margins(section, top_mm=25, bottom_mm=25, left_mm=22, right_mm=22)

    # --- Header with image (left) and sender info (right) ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.8)

    cell_img = table.cell(0, 0)
    cell_text = table.cell(0, 1)

    # Insert header image
    try:
        run = cell_img.paragraphs[0].add_run()
        run.add_picture(image_path, width=Inches(1.6))
    except Exception as e:
        cell_img.text = "[Logo aqui]"
    
    # Sender information (aligned right)
    p = cell_text.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(sender_name + "\n")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(sender_title + "\n")
    run2.italic = True
    run2.font.size = Pt(10)

    # Empty line
    doc.add_paragraph()

    # --- Date ---
    data_paragraph = doc.add_paragraph()
    data_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    data_paragraph.add_run(f"{city}, {datetime.now().strftime('%d de %B de %Y')}")

    # --- Recipient ---
    dest = doc.add_paragraph()
    dest.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dest.add_run(recipient_name).bold = True

    doc.add_paragraph()

    # --- Subject ---
    subj = doc.add_paragraph()
    subj.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = subj.add_run(subject_text)
    run.bold = True
    run.font.size = Pt(11)

    # --- Letter body ---
    chart_inserted = False
    for i, para in enumerate(body_paragraphs):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Add text with bold formatting for **text**
        add_paragraph_with_bold(p, para)

        p_format = p.paragraph_format
        p_format.space_after = Pt(8)
        p_format.line_spacing = 1.15
        
        # Insert chart after the first or second paragraph (where allocation is mentioned)
        # Look for keywords that suggest portfolio allocation discussion
        if not chart_inserted and chart_figure is not None:
            # Check if this paragraph mentions allocation, distribution, or portfolio composition
            lower_para = para.lower()
            allocation_keywords = ['alocação', 'distribuição', 'composto por', 'portfólio', 'ações', 'fundos', 'renda fixa']
            if any(keyword in lower_para for keyword in allocation_keywords) and i >= 0:
                # Insert chart after this paragraph
                doc.add_paragraph()  # Add spacing
                
                # Save chart to temporary buffer
                chart_buffer = io.BytesIO()
                chart_figure.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
                chart_buffer.seek(0)
                
                # Insert chart centered
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = chart_para.add_run()
                run.add_picture(chart_buffer, width=Inches(3.5))
                
                doc.add_paragraph()  # Add spacing after chart
                chart_inserted = True

    # If chart wasn't inserted yet, insert it after all paragraphs
    if not chart_inserted and chart_figure is not None:
        doc.add_paragraph()  # Add spacing
        
        # Save chart to temporary buffer
        chart_buffer = io.BytesIO()
        chart_figure.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
        chart_buffer.seek(0)
        
        # Insert chart centered
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = chart_para.add_run()
        run.add_picture(chart_buffer, width=Inches(3.5))
        
        doc.add_paragraph()  # Add spacing after chart

    doc.add_paragraph() 

    # --- Final paragraph before closing ---
    final_para = doc.add_paragraph()
    final_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    final_text = "Estamos à disposição para discutir mais detalhadamente os resultados e as recomendações, bem como para esclarecer quaisquer dúvidas que você possa ter. Agradecemos pela confiança depositada em nossos serviços e seguimos comprometidos em ajudá-lo a alcançar seus objetivos financeiros."
    add_paragraph_with_bold(final_para, final_text)
    
    final_para_format = final_para.paragraph_format
    final_para_format.space_after = Pt(8)
    final_para_format.line_spacing = 1.15

    doc.add_paragraph()

    # --- Closing line and signature ---
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    closing.add_run(closing_line).bold = False


    sign = doc.add_paragraph()
    sign.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sign.add_run(signature_name + "\n")
    run.bold = True
    run.font.size = Pt(11)
    run2 = sign.add_run(signature_title)
    run2.font.size = Pt(10)
    run2.italic = True

    # --- Footer (optional) ---
    sect = doc.sections[0]
    footer = sect.footer
    f_para = footer.paragraphs[0]
    f_para.text = "XP Investimentos · Av. Presidente Juscelino Kubitschek, 1909 · (11) 4003 3710"
    f_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f_para.style = doc.styles['Normal']
    f_para.paragraph_format.space_before = Pt(6)
    f_para.paragraph_format.space_after = Pt(0)
    f_para.font = f_para.runs[0].font if f_para.runs else None
    if f_para.runs:
        f_para.runs[0].font.size = Pt(9)

    # --- Save the document ---
    # Ensure output directory exists
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_file))
    print(f"Documento salvo em: {output_file}")